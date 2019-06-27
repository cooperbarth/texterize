import os, numpy as np
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.shared import RGBColor

FONT = "Courier" #any monospaced font
FONT_SIZE = 1.0
LINE_SPACING = 0.5

#writes a text block to a .docx file
def writeDoc(text_arr, chroma, write_path, overwrite):
    '''
    params:
    -text_arr: A 2-D numpy array of the text to write to the document
    -chroma: A 3-D numpy array containing the RGB values to assign to each character of text
    -write_path: The path to write the output file to
    -overwrite: A Boolean representing whether or not to overwrite an existing .docx file if found.
    '''

    document = Document()
    doc_style = document.styles['Normal']

    font = doc_style.font
    font.name = FONT
    font.size = Pt(FONT_SIZE)

    p_format = doc_style.paragraph_format
    p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_format.space_before = 0
    p_format.space_after = 0
    p_format.line_spacing = LINE_SPACING

    for i in range(text_arr.shape[0]):
        p = document.add_paragraph()
        p.style = doc_style
        for j in range(text_arr.shape[1]):
            c = p.add_run(text_arr[i][j])
            R, G, B = [int(c) for c in chroma[i][j]]
            c.font.color.rgb = RGBColor(R, G, B)
 
    if overwrite and os.path.exists(write_path):
        os.remove(write_path)
    document.save(write_path)